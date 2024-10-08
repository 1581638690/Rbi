#/bin/python
# -*- coding: utf-8 -*-


"""
doc.py
文档生成服务
#=========================


"""
import pandas as pd
import sys
import json
import shutil
from docx.shared import Inches,Pt, RGBColor,Cm,Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK,WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docxtpl import *
sys.path.append("../")
sys.path.append("./lib")
from driver.pyssdb import Client
from  avenger.fsys  import b64
from avenger.fglobals import logger,fbi_global
import os
#add by gjw on 20221104, 不再使用matplotlib库
try:
    import matplotlib.pyplot as plt
    plt.rcParams['font.family'] = 'simhei'
    plt.rcParams['axes.unicode_minus'] = False
    plt.switch_backend('agg')
except:
    plt=None

__workSpace="../workspace/"
default_img="/opt/openfbi/mPig/html/images/dbd6/chart.png"
__script_dir="script/"
sys.path.append("../")
__workSpace_report="../workspace/report/"
__html = "/opt/openfbi/report/"



"""
add by gjw on 20221010 根据报告名称生成html
"""
def word2html(df,p=""):
    
    from pydocx import PyDocX
    html = PyDocX.to_html(__workSpace_report+p+".docx")
    f = open(__html+p+".html", 'w', encoding="utf-8")
    f.write(html)
    f.close()
    return  pd.DataFrame([["/bi/report/"+p+".html"]],columns=["link"])

"""
修改doc文件，替换{{这种变量}}
@参数:
@返回:
"""
def modifiy_doc(df,p=""):
    p=p.strip()
    errMessage=[]
    status=1
    try:
        try:
            id=p.split(",")[0].strip()
            base=p.split(",")[1].strip()
            var_data=p.split(",")[2].strip()
            tbs_data=p.split(",")[3].strip()
            report_name=p.split(",")[4].strip()
        except Exception:
            errMessage.append("参数必须为5个id,base,var_data,tbs_data,report_name")
        src_path=__workSpace+"temp_word/"+id+"/"
        file_list=os.listdir(src_path)
        for file in file_list:
            if file.strip().endswith(".docx") and file.strip()!="template.docx":
                os.remove(src_path+file.strip())
        if report_name=="":
            save_path=src_path+"make_report.docx"
        else:
            save_path=src_path+report_name+".docx"
        tpl = DocxTemplate(src_path+"template.docx")
        #添加 rzc
        tpl.init_docx()
        content={}
        n=len(tpl.docx.tables)

        base_df=get_ssdb_data(base)
        vars_df=get_ssdb_data(var_data)
        tbs_df=get_ssdb_data(tbs_data)
        if base_df.index.size==0:
            if vars_df.index.size!=0:
                for index,row in vars_df.iterrows():
                    df_temp=get_ssdb_data(row["var_key"])
                    if df_temp.index.size!=0:
                        data_dic=df_temp.to_dict(orient='records')[0]
                        content[row["var"]]=data_dic
                    else:
                        errMessage.append("var_key的数据不存在")
            if tbs_df.index.size==0 and n!=0:
                for index,row in df.iterrows():
                    try:
                        tpl.docx.tables[index].autofit=True
                        tax=tpl.docx.tables[index].cell(0,0).add_paragraph('')
                        tax.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = tax.add_run("")
                        try:
                            run.add_picture(src_path+str(index)+".png",width=Inches(5.2),height=Inches(2.4))
                        except Exception:
                            run.add_picture(src_path+row["pic_name"],width=Inches(5.2),height=Inches(2.4))
                        tpl.docx.tables[index].style=tpl.docx.styles['Normal Table']
                    except Exception:
                        continue
            if tbs_df.index.size!=0:
                table_list=[]
                pics_list=[]
                for i in range(n):
                    tpl.docx.tables[i].autofit=True
                    if tpl.docx.tables[i].cell(0,0).text.strip()=="":
                        pics_list.append(tpl.docx.tables[i])
                    else:
                        table_list.append(tpl.docx.tables[i])
                if len(pics_list)!=0:
                    for i,pic_re_var in enumerate(pics_list):
                        try:
                            tax=pic_re_var.cell(0,0).add_paragraph('')
                            tax.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            run = tax.add_run("")
                            try:
                                run.add_picture(src_path+str(i)+".png",width=Inches(5.2),height=Inches(2.4))
                            except Exception:
                                run.add_picture(src_path+df.loc[i,"pic_name"],width=Inches(5.2),height=Inches(2.4))
                            tpl.docx.tables[i].style=tpl.docx.styles['Normal Table']
                        except Exception:
                            continue
                for index,row in tbs_df.iterrows():
                    table_data=get_ssdb_data(row["table_key"])
                    if row["has_header"]=="y":
                        table_data=comb_df(table_data)
                    content[row["table_var"]]=table_data.to_dict(orient='records')
                for tab_v in table_list:
                    tab_v.style=tpl.docx.styles['Table Grid']
                    tab_v.alignment=WD_TABLE_ALIGNMENT.CENTER
            tpl.render(content)
            tpl.save(save_path)
        else:
            df_single_var=get_ssdb_data(base)
            df_single_var_dict=parse_key_value(df_single_var)
            if vars_df.index.size!=0:
                for index,row in vars_df.iterrows():
                    df_temp=get_ssdb_data(row["var_key"])
                    if df_temp.index.size!=0:
                        data_dic=df_temp.to_dict(orient='records')[0]
                        df_single_var_dict[row["var"]]=data_dic
                    else:
                        errMessage.append("var_key的数据不存在")
            if tbs_df.index.size==0 and n!=0:
                for index,row in df.iterrows():
                    try:
                        tpl.docx.tables[index].autofit=True
                        tax=tpl.docx.tables[index].cell(0,0).add_paragraph('')
                        tax.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = tax.add_run("")
                        try:
                            run.add_picture(src_path+str(index)+".png",width=Inches(5.2),height=Inches(2.4))
                        except Exception:
                            run.add_picture(src_path+row["pic_name"],width=Inches(5.2),height=Inches(2.4))
                        tpl.docx.tables[index].style=tpl.docx.styles['Normal Table']
                    except Exception:
                        continue
            if tbs_df.index.size!=0:
                table_list=[]
                pics_list=[]
                for i in range(n):
                    tpl.docx.tables[i].autofit=True
                    if tpl.docx.tables[i].cell(0,0).text.strip()=="":
                        pics_list.append(tpl.docx.tables[i])
                    else:
                        table_list.append(tpl.docx.tables[i])
                if len(pics_list)!=0:
                    for i,pic_re_var in enumerate(pics_list):
                        try:
                            tax=pic_re_var.cell(0,0).add_paragraph('')
                            tax.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            run = tax.add_run("")
                            try:
                                run.add_picture(src_path+str(i)+".png",width=Inches(5.2),height=Inches(2.4))
                            except Exception:
                                run.add_picture(src_path+df.loc[i,"pic_name"],width=Inches(5.2),height=Inches(2.4))
                            tpl.docx.tables[i].style=tpl.docx.styles['Normal Table']
                        except Exception:
                            continue
                for index,row in tbs_df.iterrows():
                    table_data=get_ssdb_data(row["table_key"])
                    if row["has_header"]=="y":
                        table_data=comb_df(table_data)
                    df_single_var_dict[row["table_var"]]=table_data.to_dict(orient='records')
                for tab_v in table_list:
                    tab_v.style=tpl.docx.styles['Table Grid']
                    tab_v.alignment=WD_TABLE_ALIGNMENT.CENTER
            tpl.render(df_single_var_dict)
            tpl.save(save_path)
    except Exception as e:
        errMessage.append(str(e))
        raise Exception("报告生成失败:%s"%(e))
    if report_name+".docx" not in os.listdir(__workSpace+"temp_word/"+id):
        status=0
    else:
        if "report" not in os.listdir(__workSpace):
            os.mkdir(__workSpace+"report")
        shutil.copy(src_path+report_name+".docx",__workSpace+"report")
    if errMessage or status==0:
        return pd.DataFrame([dict(status=0,errMessage=errMessage)])
    else:
        return pd.DataFrame([dict(status=1,errMessage="")])


"""
将df表画成图像支持一下六种
line :折线图
bar : 柱状图
barh : 横向柱状图
pie :饼图
scatter :散点图
@参数:
df,p
import matplotlib
matplotlib.matplotlib_fname()
rm -rf ~/.cache/matplotlib
'/opt/fbi-base/lib/python3.7/site-packages/matplotlib-3.1.1-py3.6-linux-x86_64.egg/matplotlib/mpl-data/matplotlibrc'
注意：因为要显示中文，所以要把simhei.ttf放在/opt/fbi-base/lib/python3.7/site-packages/matplotlib/mpl-data/fonts/ttf目录下面，并执行 find -name fontList.json
./root/.cache/matplotlib/fontList.json
@返回:
"""
def generate_pic(df,p=""):

    try:

        """调用Pyecharts进行绘图"""
        py_echarts(df, p=p)

    #############结束逻辑更改###################
    except:
        """原matplotlib绘图"""
        generate_pic1(df, p=p)

"""
删除id面板对应的文件夹
"""
def dele_board(df,p=""):
    p=p.strip()
    path=__workSpace+"temp_word/"+p
    if p!="":
        os.system("rm -rf %s"%(path))
"""
通过SSDB的key获取数据
"""
def get_ssdb_data(key):#hist传入
    conn = fbi_global.get_ssdb0()
    a = conn.get(b64(key))
    if a:#如果改key存在值
        b=json.loads(a)#{"columns":["id","name"],"index":[0,1,2,3,4],"data":[["1","7"],["2","8"],["3","9"],["4","10"],["8","11"]]}
        df=pd.DataFrame(b['data'],columns=b['columns'],index=b['index'])
    else:
        df=pd.DataFrame()
    return df
"""
转化为json
"""
def get_ssdb_data1(key):
    conn= fbi_global.get_ssdb0()
    a=conn.get(b64(key))
    if a:
        b=json.loads(a)
    else:
        b={}
    return b


"""
将key,value对应的表转换为dict
"""
def parse_key_value(df):
    data={}
    for index,row in df.iterrows():
        data[row["name"]]=row["value"]
    return data
"""
将df表的列名添加到df表的第一行
"""
def comb_df(df):
    colNames=df.columns.tolist()
    df1=pd.DataFrame([colNames],columns=colNames)
    res=df1._append(df,ignore_index=True)
    return res
"""
获取面板id除了模板文件还有那些docx文件
"""
def list_report(df):
    path=__workSpace+"temp_word/"
    res=[]
    allfilelist=os.listdir(path)
    for file in allfilelist:
        filepath=os.path.join(path,file)
        for file_t in os.listdir(filepath):
            if file_t.endswith(".docx") and file_t!="template.docx":
                dic={"id":file,"report_name":file_t,"report_path":"workspace/"+"temp_word/"+file+"/"+file_t}
                res.append(dic)
    df1=pd.DataFrame(res)
    df2=pd.DataFrame([len(res)],columns=["id"])
    return df1,df2

"""
新增Pyecharts模块逻辑，一列的时候出来
"""

"""
    获取数据 Pyecharts中的数据
"""

"""调用Pyecharts进行绘图"""
def py_echarts(df,p=""):
    from pyecharts.charts import Pie, Scatter, Line, Bar,Grid
    from pyecharts import options as opts
    from snapshot_phantomjs import snapshot
    from pyecharts.render import make_snapshot
    from pyecharts.commons.utils import JsCode
    from pyecharts.options import ItemStyleOpts
    import os


    file_path = "{}/".format(os.path.dirname(os.path.abspath("/opt/openfbi/fbi-bin/static/libs/echarts.min.js")))
    #file_path = "{}/".format(os.path.dirname(os.path.abspath("/py_echarts/download/echarts.min.js")))
    
    if df.index.size == 0:
        pass
    else:

        path = __workSpace + "temp_word/" + p + "/"
            # add by gjw on 2022-0330
        import shutil, os, glob
        pngs = glob.glob(path + '*.png')
        for png in pngs:
            os.remove(png)
        for index, row in df.iterrows():
            key = row["key"].strip()
            type = row["type"].strip()
                # pic_name=row["pic_name"].strip()
            if type != "":
                # 获取到df_dict的值
                df_dict = get_ssdb_data1(key)  ##{"columns":["id","name"],"index":[0,1,2,3,4],"data":[["1","7"],["2","8"],["3","9"],["4","10"],["8","11"]]}
                if df_dict.get("index") == []:
                    if "chart.png" in os.listdir(path):
                        pass
                    else:
                        shutil.copy(default_img, path)
                    df.loc[index, "pic_name"] = "chart.png"
                    continue
                # 获取到列值
                columns_list = df_dict["columns"]  # 列名列表
                index_list = df_dict["index"]  # 索引列表
                data_list = df_dict["data"]  # 列值数组
                count = len(data_list[0])  # 获取到有几列数据
                color_list = ["rgb(65,112,156)", "rgb(76,132,183)", "rgb(85,148,204)", "rgb(122,170,218)",
                              "rgb(163,192,225)", "rgb(195,213,235)", "rgb(152,169,232)", "rgb(132,135,210)",
                              "rgb(132,173,210)", "rgb(132,210,207)"]
                #添加折线图形状列表 add by rzc 2023-04-27
                sysmbol_list=["circle","rect","roundRect","triangle","diamond"]
                #设置
                grid = Grid(init_opts=opts.InitOpts(width="1600px",height="800px"))
                if type == 'scatter':
                    if count == 1:
                        continue
                    sx = (Scatter(init_opts=opts.InitOpts(js_host=file_path, bg_color="white")))
                    sx.add_xaxis(xaxis_data=[d[0] for d in data_list])
                    sx.add_yaxis(
                            series_name="",
                            y_axis=[d[1] for d in data_list],
                            label_opts=opts.LabelOpts(is_show=True),
                        )
                    sx.set_colors([i for i in color_list])
                    sx.set_global_opts(
                        xaxis_opts=opts.AxisOpts(
                            type_="value", splitline_opts=opts.SplitLineOpts(is_show=False),
                            axislabel_opts=opts.LabelOpts(font_size=20,rotate=-15,interval=0),
                            name_rotate=60
                        ),
                        yaxis_opts=opts.AxisOpts(
                            type_="value",
                            axistick_opts=opts.AxisTickOpts(is_show=True),
                            splitline_opts=opts.SplitLineOpts(is_show=True),
                            axislabel_opts=opts.LabelOpts(font_size=20)
                        ),
                    )
                    make_snapshot(snapshot, sx.render(), path + str(index) + ".png", is_remove_html=True, pixel_ratio=1)
                elif type == 'pie':
                    y_data = [d[0] for d in data_list]
                    bx = (
                        Pie(init_opts=opts.InitOpts(js_host=file_path, bg_color="white"))
                            .add("",
                                 [list(z) for z in zip(index_list, y_data)],
                                 center=["40%", "50%"],
                                 ).set_colors(
                            [i for i in color_list]
                        ).set_global_opts(
                            legend_opts=opts.LegendOpts(orient="vertical", pos_left="right",textstyle_opts=opts.TextStyleOpts(font_size=20)),
                        ).set_series_opts(label_opts=opts.LabelOpts(formatter="{b}: {c}", font_size=20,
                                                                    ), )
                    )
                    make_snapshot(snapshot, bx.render(), path + str(index) + ".png", is_remove_html=True, pixel_ratio=1)
                elif type == "line":
                    #{"columns":["id","name"],"index":[0,1,2,3,4],"data":[["1","7"],["2","8"],["3","9"],["4","10"],["8","11"]]}
                    lx = Line(init_opts=opts.InitOpts(js_host=file_path, bg_color="white"))
                    lx.add_xaxis(xaxis_data=index_list)
                    for i in range(0, count):
                        #add by rzc 2023-04-27 添加symbol参数为折线形状列表
                        lx.add_yaxis(series_name=columns_list[i], y_axis=[y[i] for y in data_list],symbol=sysmbol_list[i],symbol_size=7)
                    lx.set_colors([i for i in color_list])
                    lx.set_global_opts(
                        xaxis_opts=opts.AxisOpts(
                            axislabel_opts=opts.LabelOpts(font_size=20,rotate=-15,interval=0),name_rotate=60),
                        yaxis_opts=opts.AxisOpts(
                            type_="value",
                            splitline_opts=opts.SplitLineOpts(is_show=True)),
                    )
                    make_snapshot(snapshot, lx.render(), path + str(index) + ".png", is_remove_html=True, pixel_ratio=1)
                elif type == "bar":
                    bx = (Bar(init_opts=opts.InitOpts(js_host=file_path, bg_color="white")))
                    bx.add_xaxis(xaxis_data=index_list)
                    #add by rzc 2023-04-15 单条数据修改为颜色不同
                    if count ==1:

                        bx.add_yaxis(series_name=columns_list[0], y_axis=[y[0] for y in data_list], itemstyle_opts=ItemStyleOpts(color=JsCode("""
                                function(params) {
                                    var colorList = ['rgb(65,112,156)', 'rgb(76,132,183)', 'rgb(85,148,204)', 'rgb(122,170,218)',
                                                      'rgb(163,192,225)', 'rgb(195,213,235)', 'rgb(152,169,232)', 'rgb(132,135,210)',
                                                      'rgb(132,173,210)', 'rgb(132,210,207)'];
                                    return colorList[params.dataIndex];
                                }
                            """)),bar_width="30%")
                    else:
                        for i in range(0, count):
                            bx.add_yaxis(series_name=columns_list[i], y_axis=[y[i] for y in data_list],color=color_list[i],bar_width="30%")
                    bx.set_colors([i for i in color_list])
                    bx.set_series_opts(
                        markline_opts=opts.MarkLineOpts(
                            data=[
                                opts.MarkLineItem(type_="average", name="平均值")
                            ],
                            linestyle_opts=opts.LineStyleOpts(color="green",type_="dotted")
                        ),
                    )
                    bx.set_global_opts(xaxis_opts=opts.AxisOpts(
                        axislabel_opts=opts.LabelOpts(font_size=20,rotate=-15,interval=0),name_rotate=60),

                    )
                    grid.add(bx, grid_opts=opts.GridOpts(pos_bottom="50%",pos_right="20%",pos_left="20%"))
                    make_snapshot(snapshot, bx.render(), path + str(index) + ".png", is_remove_html=True, pixel_ratio=1)
                elif type == "barh":
                    bar = Bar(init_opts=opts.InitOpts(js_host=file_path, bg_color="white"))
                    bar.add_xaxis(xaxis_data=index_list)
                    ##add by rzc 2023-04-15 单条数据修改为颜色不同
                    if count==1:
                        bar.add_yaxis(series_name=columns_list[0], y_axis=[y[0] for y in data_list],
                                     itemstyle_opts=ItemStyleOpts(color=JsCode("""
                                                        function(params) {
                                                            var colorList = ['rgb(65,112,156)', 'rgb(76,132,183)', 'rgb(85,148,204)', 'rgb(122,170,218)',
                                                                              'rgb(163,192,225)', 'rgb(195,213,235)', 'rgb(152,169,232)', 'rgb(132,135,210)',
                                                                              'rgb(132,173,210)', 'rgb(132,210,207)'];
                                                            return colorList[params.dataIndex];
                                                        }
                                                    """)), bar_width="30%")
                    else:
                        for i in range(0, count):
                            bar.add_yaxis(series_name=columns_list[i], y_axis=[y[i] for y in data_list],bar_width="30%")
                    bar.reversal_axis()
                    bar.set_colors([i for i in color_list])
                    bar.set_series_opts(label_opts=opts.LabelOpts(position="right"),  # 标签字体调整
                                        markline_opts=opts.MarkLineOpts(
                                            data=[
                                                opts.MarkLineItem(type_="average", name="平均值")
                                            ],
                                            linestyle_opts=opts.LineStyleOpts(color="green",type_="dotted")
                                        ))
                    bar.set_global_opts(
                        yaxis_opts=opts.AxisOpts(
                            axislabel_opts=opts.LabelOpts(font_size=20,rotate=-15),name_rotate=60), )

                    make_snapshot(snapshot, bar.render(), path + str(index) + ".png", is_remove_html=True,
                                  pixel_ratio=1)
                elif type == "hist":
                    bx = Bar(init_opts=opts.InitOpts(js_host=file_path, bg_color="white"))
                    bx.add_xaxis(xaxis_data=index_list)
                    for i in range(0, count):
                        bx.add_yaxis(series_name=columns_list[i], y_axis=[y[i] for y in data_list],gap="0%",
                                     category_gap=0)
                    bx.set_colors([i for i in color_list])
                    bx.set_series_opts(
                        label_opts=opts.LabelOpts(is_show=False),
                        markline_opts=opts.MarkLineOpts(
                            data=[
                                opts.MarkLineItem(type_="average", name="平均值")
                            ],
                            linestyle_opts=opts.LineStyleOpts(color="green",type_="dotted")
                        ))
                    bx.set_global_opts(
                        xaxis_opts=opts.AxisOpts(
                            axislabel_opts=opts.LabelOpts(font_size=20,rotate=-15,interval=0),name_rotate=60), )
                    grid.add(bx, grid_opts=opts.GridOpts(pos_bottom="50%",pos_right="20%",pos_left="20%"))
                    make_snapshot(snapshot, grid.render(), path + str(index) + ".png", is_remove_html=True, pixel_ratio=1)


"""原matplotlib绘图"""

def generate_pic1(df,p=""):
    
    if df.index.size == 0:
        pass
    else:
        try:
            path = __workSpace + "temp_word/" + p + "/"
            # add by gjw on 2022-0330
            import shutil, os, glob
            pngs = glob.glob(path + '*.png')
            for png in pngs:
                os.remove(png)
            for index, row in df.iterrows():
                key = row["key"].strip()
                type = row["type"].strip()
                # pic_name=row["pic_name"].strip()
                if type != "":
                    df_temp = get_ssdb_data(
                        key)  # key如果等于 pie  #df_temp=pd.DataFrame(b['data'],columns=b['columns'],index=b['index'])
                    if df_temp.index.size == 0:
                        # if "%d.png"%(index) in os.listdir(path):
                        #	os.remove(path+"%d.png"%(index))
                        if "chart.png" in os.listdir(path):
                            pass
                        else:
                            shutil.copy(default_img, path)
                        df.loc[index, "pic_name"] = "chart.png"
                        continue
                    if type == 'scatter':
                        ax = df_temp.plot(x=df_temp.columns.tolist()[0], y=df_temp.columns.tolist()[1],kind='scatter', figsize=(8.6, 4), rot=0)
                    elif type == 'pie':
                        ax = df_temp.plot(kind='pie', y=df_temp.columns.tolist()[0], figsize=(8.6, 4), rot=0)
                        ax.axis('equal')
                        ax.legend(bbox_to_anchor=(-0.1, 1.0), loc='upper left')
                    else:
                        ax = df_temp.plot(kind=type, figsize=(8.6, 4), rot=0)
                    fig = ax.get_figure()
                    fig.savefig(path + str(index) + ".png")
        except Exception as e:
            raise Exception(e)
